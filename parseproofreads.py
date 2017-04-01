#!/usr/bin/python
# -*- coding: utf-8 -*-
#
### -*- coding: latin-1 -*-
# srchan 2016.05.10
# revised 2016.11.29 srchan - to look at DOCX files, output CSV

# parseproofreads.py - for UC Davis Virtual Collaborative Care Network project (BHCE Translate)
# Parser for proofreadings from interpreters.
# 
# Outputs the following, given a file:
# - total # of tokens
# - additions
# - deletions
# - unusual changes, preserved meaning
# - drastic changes, meaning NOT preserved
# 
# ... per word
# ... per sentence (Delimited by period, question mark, exclamation mark.) (extra credit)
# Configure by changing the colors at the top of the document.


####### LIBRARY IMPORTS
# Uses Python-DocX
# https://python-docx.readthedocs.io/en/latest/
# https://pypi.python.org/pypi/python-docx

from docx import Document
from collections import defaultdict
import getopt
import csv
import os
import logging
import sys
import time

# _debug = False  # Show debugging messages
# logging.basicConfig(level=logging.WARNING)
logging.basicConfig(format='[%(asctime)s][%(name)s] %(levelname)s :: %(message)s')
log = logging.getLogger(__name__)
log.setLevel(logging.WARNING)


def usage(scriptname='<scriptname>'):
	print """
USAGE
"""
	print "   " + str(scriptname) + " -i interpreter.docx -e engines.docx -o results.csv"
	print """

REQUIRED
-i FILENAME, --interpreterdocx=FILENAME
	This docx file must contain 1 table, from the interpreters.
	The interpreters must 
-e FILENAME, --enginekeydocx=FILENAME
	This docx file must contain 1 table of the same size as specified under "interpreterdocx".
	Each cell with the engine name corresponds to the "interpreterdocx" table.
-o FILENAME, --outputcsv=FILENAME
	This is the name of the CSV file to output.
	The structure of the CSV file will appear like so:
	            total words     total added words      total deleted words      total unusual words     total drastic words
   	engine1
   	engine2
   	engine3

OPTIONAL
-h, --help
	Shows this help



	"""

def main(argv, scriptname):                         
	try:                                
		opts, args = getopt.getopt(argv, "hi:e:o:d", ["help", "interpreterdocx=", "enginekeydocx=", "outputcsv=", ""])
	except getopt.GetoptError:
		usage(scriptname)
		sys.exit(2)


	for opt, arg in opts:		# Process the command-line arguments.
		if opt in ('-h','--help'):
			usage(scriptname)
			sys.exit()
		elif opt == '-d':
			#global _debug
			#_debug = True
			log.setLevel(logging.DEBUG)
		elif opt in ('-i','interpreterdocx'):
			if os.path.isfile(arg):
				interpreter_docx = arg
			else:
				raise IOError ("Interpretation DOCX file does not exist.")
		elif opt in ('-e','enginekeydocx'):
			if os.path.isfile(arg):
				enginekey_docx = arg
			else:
				raise IOError ("Engine DOCX file does not exist.")
		elif opt in ('-o','outputcsv'):
			output_csv = arg
			# This file does not necessarily need to exist.
		else:
			pass

	
	# Attempt to open each file.
	# If interpreterdocx, enginekeydocx, outputcsv do not exist,
	# then show error message.
	try:
		interpretation_table = getTable(Document(interpreter_docx))
	except NameError:
		raise NameError ("Argument for interpreterdocx needs to be specified.")

	try:
		enginekey_table = getTable(Document(enginekey_docx))
	except NameError:
		raise NameError ("Argument for enginekeydocx needs to be specified.")

	try:
		output_csv   # This just tests if variable is set.
	except NameError:
		raise NameError ("Argument for outputcsv needs to be specified, so that a CSV file with the statistical results can be generated.")


	# Initialize the total counts we want.
	# The defaultdict basically allows us to ask it for keys
	# that may not exist, but will always return 0.
	# For instance, if total_counts['Bluemix'] doesn't exist,
	# it will return 0.
	# Each KEY is an engine name.
	# Each VALUE is the counts associated with that engine.
	total_counts = defaultdict(newCountDictionary) # uses the function newCountDictionary as a lambda to return a new dictionary
	# total_word_count = defaultdict(lambda: 0)
	# total_added_word_count = defaultdict(lambda: 0)
	# total_deleted_word_count = defaultdict(lambda: 0)  # FUTURE
	# total_unusual_word_count = defaultdict(lambda: 0)
	# total_drastic_word_count = defaultdict(lambda: 0)
	# total_paragraph_count = defaultdict(lambda: 0)
	# total_sentence_count = defaultdict(lambda: 0)

	# Initialize a set (basically a list that holds unique items with no repeats)
	# -- that is, the name of our engines.
	# This is so we can iterate later on when we write to CSV.
	all_engine_names  = set()

	# Check to make sure the length of rows, columns are the same.
	if ((len(interpretation_table.rows) != len(enginekey_table.rows)) 
		or (len(interpretation_table.columns) != len(enginekey_table.columns))):
		raise IOError ("The size of the tables for " + str(interpreter_docx) + " and " + enginekey_docx + " do not match.")


 
	log.info("Starting job at " + time.strftime("%c"))
 

	# For each row in the table, though
	# skip row 0 -- as this is a header. (So we start at 1.)
	for i in range(1, len(interpretation_table.rows)):
		# For each column in the table,...
		# Skip column 0 -- as this is the original interpretation.
		for j in range(1, len(interpretation_table.columns)):
			engine_name = enginekey_table.cell(i,j).text
			log.debug(u'Got engine name ' + engine_name)
			# all_engine_names = all_engine_names.add(engine_name)
			cell_paragraphs = interpretation_table.cell(i,j).paragraphs
			cell_statistics = processCell (cell_paragraphs)

			for statistic_name, value in cell_statistics.items():
					total_counts[engine_name][statistic_name] += value
			# total_word_count[engine_name] += 
			# total_added_word_count[engine_name] += 
			# total_deleted_word_count[engine_name] +=   # FUTURE
			# total_unusual_word_count[engine_name] += 
			# total_drastic_word_count[engine_name] += 
			# total_paragraph_count[engine_name] += 
			# total_sentence_count[engine_name] += 




	# Write out the CSV file.
	with open(output_csv, 'wb') as csvfile:
		w = csv.writer(csvfile, dialect='excel')

		# Get the keys for newCountDictionary(), sorted / alphabetized.
		# i.e. ['added_word_count', 'deleted_word_count', 'drastic_word_count', 'paragraph_count', 'sentence_count', 'unusual_word_count', 'word_count']
		alphabetized_counter_names = sorted(newCountDictionary().keys())

		# This is where you can remove certain keys that you don't want reported
		# in the CSV file.
		# FUTURE: Make this more elegant by not including them in the first place. :-)
		# << none excluded at this time >>


		# Write out the header: all of the counts.
		# i.e. [BLANK, 'added_word_count', 'deleted_word_count', 'drastic_word_count' ...]
		w.writerow([''] + alphabetized_counter_names)

		# Write the statistics to the CSV file, for each engine.
		for engine_name, engine_statistics in total_counts.items():
			row = [str(engine_name)]

			# We have to iterate over each item in counter_names, to ensure fidelity...
			# ...in the very off-case that such an item doesn't exist anyways.
			for key in alphabetized_counter_names:
				row.append(str(engine_statistics[key]))

			# Our final row to write to the CSV looks like: [engine_name, 2, 5, 10, 1, ...]
			w.writerow(row)


	log.info("Finished job at " + time.strftime("%c"))



# Returns a new dictionary that holds counts of words
def newCountDictionary ():
	return {'word_count':0,
			'added_word_count':0,
			'deleted_word_count':0,  # FUTURE
			'unusual_word_count':0,
			'drastic_word_count':0,
			'paragraph_count':0,
			'sentence_count':0 }

# Given a document (DOCX) file,
# returns the paragraphs from the single table.
def getTable (document):
	if (len(document.tables) == 1): 
		return document.tables[0]
	else:
		raise KeyError("Needs to have exactly one table within document.")

# This is used by processCell().
# This adds a word count to #1, with or without #2:
# (1) the total word count, and
# (2) a particular tagged count
# ...to the word_counts dictionary.
# By default it returns True once it's executed.
def add_to_word_count(word_counts, is_on_added_tag=False, is_on_unusual_tag=False, is_on_drastic_tag=False):		
	word_counts['word_count'] += 1 # added to # of words
	
	if is_on_added_tag:
		word_counts['added_word_count'] += 1
	if is_on_unusual_tag:
		word_counts['unusual_word_count'] += 1
	if is_on_drastic_tag:
		word_counts['drastic_word_count'] += 1
	# TODO: deleted word / strikethrough
	return True


# This is used by processCell().
# This adds to the sentence_count.
# By default it returns True once it's executed.
def add_to_sentence_count(word_counts):
	word_counts['sentence_count'] += 1
	return True


# This is used by processCell().
# This adds to deleted word count.
# By default it returns True once it's executed.
def add_to_deletions(word_counts):
	log.debug('Found X deletion.')
	word_counts['deleted_word_count'] += 1
	return True

# Given a series of paragraphs (see Document module),
# returns a dictionary of statistics for that.
def processCell (interpretation_paragraphs):
	results = newCountDictionary()

	for paragraph in interpretation_paragraphs:
		results['paragraph_count'] += 1

		is_on_added_tag = False     # No words have been processed yet.
		is_on_deleted_tag = False   # FUTURE
		is_on_unusual_tag = False
		is_on_drastic_tag = False
		is_on_a_word = False

		# For each run within the current "paragraph...""
		for j in range(0, len(paragraph.runs)):

			# What if we start with a space?
			was_on_added_tag = is_on_added_tag
			was_on_deleted_tag = is_on_deleted_tag   # FUTURE
			was_on_unusual_tag = is_on_unusual_tag
			was_on_drastic_tag = is_on_drastic_tag
			was_on_a_word = is_on_a_word

			run = paragraph.runs[j]
			log.debug(u"Run " + run.text)
			is_on_added_tag = not ((run.underline == False) or (run.underline == None))
			is_on_deleted_tag = False # We will count the number of X's in the later FOR loop when we process each character.
				# run.text.count('X')  # This code counts how many characters is an X.
				# TODO: We may not need "is_on_deleted_tag"
			# TODO: is there a strikethrough?
			try:
				is_on_unusual_tag = (run.font.color.rgb[1] > 120)  # Any bright green used ([0, 1, 2] == R G B).
				is_on_drastic_tag = (run.font.color.rgb[0] > 120)  # Any bright red used ([0, 1, 2] == R G B).
				# TODO: is there a blue highlight?
			except TypeError:
				is_on_unusual_tag = False
				is_on_drastic_tag = False
				# log.debug( run.text.encode('utf-8') + " color is " + str(run.font.color) + " of type " + str(run.font.color.type) )
				# TODO: Pay attention if this error does get triggered.
				# TODO: Change print to an error log.
			is_on_a_word = False


			#if (is_on_unusual_tag):
			#	print ('"' + str(run.text) + '"')


			for i in range(0, len(run.text)): # For each character in the run...
				character = run.text[i]
				at_deleted_item = character in (u'X', 'X')  # We're safely assuming that X really isn't used anywhere else in these transcripts...
				at_end_of_word = character in (" ", "\n", "\r\n", "\r") #delimiters
				at_end_of_sentence = character in ("?", ".", "!")
				at_end_of_run = (i == len(run.text) - 1)   # Have we reached the end?


				at_end_of_paragraph = at_end_of_run & (j == len(paragraph.runs) - 1)

				# Any time an X is encountered, count that as one item removed.
				if at_deleted_item:
					add_to_deletions(results)

				# TODO: case where a run is just one word: run "encu", then run "entra", then " ".
				# This happens if we're right at the start.
				if ((i == 0) & (at_end_of_word | at_end_of_sentence) & was_on_a_word):
					add_to_word_count(results, was_on_added_tag, was_on_unusual_tag, was_on_drastic_tag)  # NOTE: this is "WAS" on tag, not "IS" on tag
					was_on_a_word = False

				# Count the number of words and sentences.
				elif at_end_of_paragraph:
					# added to # of words — because there is no delimiter at end
					add_to_word_count(results, is_on_added_tag, is_on_unusual_tag, is_on_drastic_tag)  # NOTE: this is "IS" on tag, not "WAS" on tag
					add_to_sentence_count(results)
					is_on_a_word = False
					#if is_on_unusual_tag:
					#	print (" ......")

				elif at_end_of_sentence: 
					# TODO: I wonder why it's not counted here...
					add_to_sentence_count(results)
					is_on_a_word = False
					#if is_on_unusual_tag:
					#	print (" ...")

				elif at_end_of_word:
					add_to_word_count(results, is_on_added_tag, is_on_unusual_tag, is_on_drastic_tag)  # NOTE: this is "IS" on tag, not "WAS" on tag
					is_on_a_word = False

					#if is_on_unusual_tag:
					#	print (" .")

				else:
					is_on_a_word = True


			# if len(run.text) > 0:

	
	return results



if __name__ == "__main__":
	main(sys.argv[1:], sys.argv[0])






