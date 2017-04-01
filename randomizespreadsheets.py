#!/usr/bin/python
# -*- coding: utf-8 -*-
#
#
# srchan 2016.06.23
# requires Python 2.7

# revising 2017.02.14 to include:
# - docx export support
# - creating a spreadsheet of engine names


# We need a script that will randomize the order of cells
# INPUT: CSV files of each separate translation engine
# OUTPUT: 
## - two CSV file (one with translations, one with engine names)
## - two DOCX files

import sys  
reload(sys)  
sys.setdefaultencoding('utf8')


import csv
import os
import random
import docx


## VARIABLES
rootdir = '.'

engines = {} 

resultTranslationCSVFilename = "result_translation.csv"
resultEnginenameCSVFilename = "result_enginename.csv"
resultTranslationDOCXFilename = "result_translation.docx"
resultEnginenameDOCXFilename = "result_enginename.docx"

resultFilenameList = (resultTranslationCSVFilename, resultEnginenameCSVFilename, resultTranslationDOCXFilename, resultEnginenameDOCXFilename)


print '----------'
print '----------'
print 'Randomize order of translations'


for subdir, dirs, files in os.walk(rootdir):
	print '----------'
	print 'Subdirectory: ' + str(subdir)
	# print files



	#original_rows = []
	#actual_translation_rows = []
	#test_translation_rows = []

	rows = []
	engineNames = set([])

	for filename in files: # os.listdir(subdir):
		if not filename.endswith('.csv'):
			continue # skip non-csv file

		if filename in resultFilenameList:
			continue # skip as this is an actual result file


		csvFilename = os.path.join(subdir, filename)
		engine = filename.lower().replace('.csv','')
		engineNames.add(engine)

		# csvRows = []
		# bleuScores = []

		currentFile = open(csvFilename, 'rU')
		currentReader = csv.reader(currentFile, dialect=csv.excel)
		print ('-- Filename :' + filename)
		row_number = 0
		for row in currentReader:
			#if currentReader.line_num == 1:
			#	continue  #skip first row


			try:
				translated_text = (engine,row[2])
			except IndexError:
				translated_text = (engine,"")

			try:
				actual_text = row[1]
			except IndexError:
				actual_text = ""


			if (row_number >= len(rows)):  # IF the row doesn't exist yet
				try:
					original_text = row[0]
				except IndexError:
					original_text = ""



				rows.append({'original':original_text, 'actual':actual_text, 'translated':[translated_text, ('actual', actual_text)]})
				# 2016.06.28 - modify the previous row: added the tuple containing actual_text.

			else:
				try:
					rows[row_number]['translated'].append(translated_text)
					# rows[row_number]['translated'].append(('actual', actual_text))
				except IndexError:
					print ("Failed to insert new translation at row " + str(row_number))




			row_number = row_number + 1


		currentFile.close()


	# Create new files. We will create both CSV files, and Word versions of the files.
	# There will be a total of 4 files created.
	resultTranslationCSVFilepath = os.path.join(subdir, resultTranslationCSVFilename)
	resultEnginenameCSVFilepath = os.path.join(subdir, resultEnginenameCSVFilename)
	resultTranslationDOCXFilepath = os.path.join(subdir, resultTranslationDOCXFilename)
	resultEnginenameDOCXFilepath = os.path.join(subdir, resultEnginenameDOCXFilename)


	with open(resultTranslationCSVFilepath,'wb') as resultTranslationCSVFile, open(resultEnginenameCSVFilepath,'wb') as resultEnginenameCSVFile:
		resultTranslationCSVWriter = csv.writer(resultTranslationCSVFile)
		resultEnginenameCSVWriter = csv.writer(resultEnginenameCSVFile)
		resultTranslationDOCX = docx.Document()
		resultEnginenameDOCX = docx.Document()
		
		resultTranslationDOCXTable = resultTranslationDOCX.add_table(rows=0, cols=2+len(engineNames)) 
		resultEnginenameDOCXTable = resultEnginenameDOCX.add_table(rows=0, cols=2+len(engineNames)) 

		resultTranslationDOCXTable.style = "Light Shading" 
		resultEnginenameDOCXTable.style = "Light Shading Accent 1"  # 1 - red, # 2 - blue

		print "Columns in table " + str(len(engineNames))


		for i in range (0, len(rows)):
			row = rows[i]
			random.shuffle(row['translated'])   # This is where randomisation occurs!



			# CONSTRUCT THE ROW
			#try:
			newTranslationRow = [row['original']]   # Changed 2017.02.13 to omit: , row['actual']
			newEngineRow = [row['original']]

			if i == 0: # Is this the header row?
				headerRow = ["Translation " + str(n) for n in range(1,len(engineNames) + 2)]

				newEngineRow.extend(headerRow)   # the +1 is to include the original translation
				newTranslationRow.extend(headerRow)
			else:
				for translatedTuple in row['translated']:
					newEngineRow.append(translatedTuple[0])  # key value: name of engine
					newTranslationRow.append(translatedTuple[1])  # data value: actual translation
			
			# ADD THE ROW TO CSV
			resultEnginenameCSVWriter.writerow(newEngineRow)
			resultTranslationCSVWriter.writerow(newTranslationRow)


			# ADD THE ROW TO DOCX -- a bit more complicated and requires iterating...
			translationDOCXRow = resultTranslationDOCXTable.add_row().cells
			enginenameDOCXRow = resultEnginenameDOCXTable.add_row().cells

			for j in range (0, len(newTranslationRow)):
			#	print "Inserting translationDOCXRow " + unicode(newTranslationRow[j]) + " i=" + str(i) + " j=" + str(j)
				translationDOCXRow[j].text = unicode(newTranslationRow[j])
			for j in range (0, len(newEngineRow)):
			#	print "Inserting enginenameDOCXRow " + unicode(newEngineRow[j]) + " i=" + str(i) + " j=" + str(j)
				enginenameDOCXRow[j].text = unicode(newEngineRow[j])

			#except IndexError:
			#	print ("Failed to write row " + str(row))
		resultTranslationCSVFile.close()
		resultEnginenameCSVFile.close()

		resultTranslationDOCX.save(resultTranslationDOCXFilepath)
		resultEnginenameDOCX.save(resultEnginenameDOCXFilepath)



# Things to write again:
# - Write the CSV files with separate files.
# - Write it so that each XLSX file is created.